from docx import Document


def create_doc(dict_output, titles, kenteken):
    i = 0
    filename = str(kenteken) + '_api.docx'
    document = Document()
    for api_response in dict_output:
        title = titles[i].title() + ' api_response'
        document.add_heading(title, 0)
        table = document.add_table(rows=1, cols=2)
        for question in dict_output[api_response]:
            row_cells = table.add_row().cells
            row_cells[0].text = question
            row_cells[1].text = dict_output[api_response][question]
        i += 1
        document.add_page_break()

    document.save(filename)
